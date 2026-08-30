"""Document extraction (v32) — documents in, structured data out.

The ingestion half of "Document AI": turn PDFs, scans, Office files and
plain text into (a) full text and (b) TABLES that can become datasets.
Extracted tables flow into the v27 dataset store, so anything a document
contains becomes SQL-queryable, app-buildable and dashboard-able.

Engines (all local, no external calls):
  pdf   — pdfplumber: text per page + ruled table detection (lattice)
  ocr   — pytesseract over images (png/jpg/bmp/tiff/webp)
  docx  — python-docx: paragraphs + tables
  xlsx  — openpyxl: one table per sheet (sheets act as "pages")
  csv/tsv / json / txt/md — decode + parse

Every table row is a list of cleaned string cells (``_clean_cell``); the
API layer turns the BEST table into dataset rows via ``table_to_items``.
Anything unsupported raises ValueError with an end-user message — the API
maps that to 400/415 and the engine node surfaces it as a run error.
"""

from __future__ import annotations

import csv
import io
import json
import os
import re

MAX_TEXT_CHARS = 200_000
MAX_TABLES = 50
MAX_TABLE_ROWS = 10_000

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
SUPPORTED_EXTS = sorted(
    {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".tsv", ".json", ".txt", ".md"} | IMAGE_EXTS
)


def supported(filename: str) -> bool:
    return os.path.splitext((filename or "").lower())[1] in SUPPORTED_EXTS


# ----------------------------------------------------------------- cells
def _clean_cell(v) -> str:
    """Any engine cell → presentable string (21.0 → "21", None → "")."""
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    if isinstance(v, str):
        return v.strip()
    return str(v).strip()


def _clean_row(row) -> list[str]:
    return [_clean_cell(c) for c in row]


def _non_empty_rows(rows: list[list[str]]) -> list[list[str]]:
    return [r for r in rows if any(c for c in r)]


# ----------------------------------------------------------------- engines
def _extract_pdf(content: bytes) -> dict:
    import pdfplumber

    text_parts: list[str] = []
    tables: list[dict] = []
    pages = 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            pages += 1
            try:
                page_text = page.extract_text() or ""
            except Exception:  # noqa: BLE001 — a broken page must not kill the doc
                page_text = ""
            if page_text:
                text_parts.append(page_text)
            if len(tables) < MAX_TABLES:
                try:
                    raw_tables = page.extract_tables() or []
                except Exception:  # noqa: BLE001
                    raw_tables = []
                for t in raw_tables:
                    rows = _non_empty_rows([_clean_row(r) for r in t][:MAX_TABLE_ROWS])
                    if rows:
                        tables.append(
                            {"page": pno, "rows": rows, "n_rows": len(rows), "n_cols": max(len(r) for r in rows)}
                        )
                    if len(tables) >= MAX_TABLES:
                        break
    return {
        "engine": "pdf",
        "pages": pages,
        "text": "\n\n".join(text_parts)[:MAX_TEXT_CHARS],
        "tables": tables,
        "meta": {},
    }


def _extract_image(content: bytes) -> dict:
    import pytesseract
    from PIL import Image

    try:
        img = Image.open(io.BytesIO(content)).convert("RGB")
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read image: {exc}") from exc
    try:
        text = pytesseract.image_to_string(img) or ""
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"OCR failed (is the tesseract binary installed?): {exc}") from exc
    return {
        "engine": "ocr",
        "pages": 1,
        "text": text.strip()[:MAX_TEXT_CHARS],
        "tables": [],
        "meta": {},
    }


def _extract_docx(content: bytes) -> dict:
    import docx as docx_lib

    try:
        doc = docx_lib.Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read Word document: {exc}") from exc
    text = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
    tables: list[dict] = []
    for t in doc.tables[:MAX_TABLES]:
        rows = _non_empty_rows([[_clean_cell(c.text) for c in r.cells] for r in t.rows][:MAX_TABLE_ROWS])
        if rows:
            tables.append(
                {"page": len(tables) + 1, "rows": rows, "n_rows": len(rows), "n_cols": max(len(r) for r in rows)}
            )
    return {
        "engine": "docx",
        "pages": 1,
        "text": text[:MAX_TEXT_CHARS],
        "tables": tables,
        "meta": {},
    }


def _extract_xlsx(content: bytes) -> dict:
    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Could not read workbook: {exc}") from exc
    sheets = list(wb.sheetnames)
    text_parts: list[str] = []
    tables: list[dict] = []
    for sno, name in enumerate(sheets[:5], start=1):
        ws = wb[name]
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            rows.append(_clean_row(row))
            if len(rows) >= MAX_TABLE_ROWS:
                break
        rows = _non_empty_rows(rows)
        if rows:
            tables.append({"page": sno, "rows": rows, "n_rows": len(rows), "n_cols": max(len(r) for r in rows)})
            text_parts.append(f"[sheet: {name}] " + " | ".join(rows[0]))
    wb.close()
    return {
        "engine": "xlsx",
        "pages": len(sheets),
        "text": "\n".join(text_parts)[:MAX_TEXT_CHARS],
        "tables": tables,
        "meta": {"sheets": sheets},
    }


def _decode(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _extract_csv(content: bytes, delimiter: str) -> dict:
    text = _decode(content)
    rows = _non_empty_rows([_clean_row(r) for r in csv.reader(io.StringIO(text), delimiter=delimiter)][:MAX_TABLE_ROWS])
    table = None
    if rows:
        table = {"page": 1, "rows": rows, "n_rows": len(rows), "n_cols": max(len(r) for r in rows)}
    return {
        "engine": "csv",
        "pages": 1,
        "text": text[:MAX_TEXT_CHARS],
        "tables": [table] if table else [],
        "meta": {},
    }


def _json_cell(v) -> str:
    if isinstance(v, (dict, list)):
        return json.dumps(v, ensure_ascii=False, default=str)
    return _clean_cell(v)


def _extract_json(content: bytes) -> dict:
    text = _decode(content)
    try:
        payload = json.loads(text)
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"Invalid JSON: {exc}") from exc
    tables: list[dict] = []
    items = payload if isinstance(payload, list) else [payload]
    if items and all(isinstance(i, dict) for i in items):
        header: list[str] = []
        for item in items:
            for k in item.keys():
                if k not in header:
                    header.append(k)
        rows = [header] + [[_json_cell(item.get(k)) for k in header] for item in items[:MAX_TABLE_ROWS]]
        tables.append({"page": 1, "rows": rows, "n_rows": len(rows), "n_cols": len(header)})
    return {
        "engine": "json",
        "pages": 1,
        "text": text[:MAX_TEXT_CHARS],
        "tables": tables,
        "meta": {},
    }


def _extract_text(content: bytes) -> dict:
    return {
        "engine": "text",
        "pages": 1,
        "text": _decode(content)[:MAX_TEXT_CHARS],
        "tables": [],
        "meta": {},
    }


# ----------------------------------------------------------------- facade
def extract_document(filename: str, content: bytes) -> dict:
    """File bytes → {engine, pages, chars, text, tables, meta}. ValueError on unsupported/broken input."""
    ext = os.path.splitext((filename or "").lower())[1]
    if ext not in SUPPORTED_EXTS:
        raise ValueError(
            f"Unsupported document type {ext or '(none)'} — supported: {', '.join(SUPPORTED_EXTS)}"
        )
    if not content:
        raise ValueError("The document is empty")
    if ext == ".pdf":
        result = _extract_pdf(content)
    elif ext in IMAGE_EXTS:
        result = _extract_image(content)
    elif ext == ".docx":
        result = _extract_docx(content)
    elif ext in (".xlsx", ".xls"):
        result = _extract_xlsx(content)
    elif ext == ".csv":
        result = _extract_csv(content, ",")
    elif ext == ".tsv":
        result = _extract_csv(content, "\t")
    elif ext == ".json":
        result = _extract_json(content)
    else:
        result = _extract_text(content)
    result["chars"] = len(result.get("text", ""))
    return result


def best_table(result: dict) -> dict | None:
    """The extractable table with the most rows (needs header + ≥1 data row)."""
    candidates = [t for t in result.get("tables", []) if t.get("n_rows", 0) >= 2]
    if not candidates:
        return None
    return max(candidates, key=lambda t: (t["n_rows"], -t.get("page", 0)))


def table_to_items(table: dict) -> list[dict]:
    """Table rows → objects keyed by the header row (pad short rows, drop extras)."""
    rows = table["rows"]
    header = rows[0]
    names: list[str] = []
    for i, h in enumerate(header):
        name = h or f"col_{i + 1}"
        base, n = name, 2
        while name in names:
            name = f"{base}_{n}"
            n += 1
        names.append(name)
    items: list[dict] = []
    for r in rows[1:]:
        padded = list(r) + [""] * (len(names) - len(r))
        items.append(dict(zip(names, padded[: len(names)])))
    return items


# ----------------------------------------------------------------- typing
_INT_RE = re.compile(r"-?(0|[1-9][0-9]*)\Z")
_FLOAT_RE = re.compile(r"-?([0-9]+\.[0-9]*|\.[0-9]+)([eE][+-]?[0-9]+)?\Z")


def coerce_items_dtypes(items: list[dict]) -> list[dict]:
    """Numeric-looking string columns become real numbers.

    Extraction cells are strings (OCR/PDF give text); a column whose
    non-empty values ALL parse as int/float is cast so datasets get real
    dtypes (qty → integer, amount → number). Leading-zero strings ("007")
    and anything with separators stay text; mixed columns stay text.
    """
    if not items:
        return items
    keys = list(items[0].keys())
    castable: dict[str, callable] = {}
    for k in keys:
        vals = [str(i.get(k, "")).strip() for i in items]
        non_empty = [v for v in vals if v != ""]
        if not non_empty:
            continue
        if all(_INT_RE.fullmatch(v) for v in non_empty):
            castable[k] = int
        elif all(_FLOAT_RE.fullmatch(v) for v in non_empty):
            castable[k] = float
    out: list[dict] = []
    for item in items:
        row = {}
        for k in keys:
            v = item.get(k)
            if k in castable and isinstance(v, str) and v.strip() != "":
                row[k] = castable[k](v.strip())
            else:
                row[k] = v
        out.append(row)
    return out
