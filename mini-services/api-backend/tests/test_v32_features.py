"""V32 feature tests: Document AI - documents in, structured data out.

* /documents/engines: capability probe (OCR availability via tesseract binary,
  supported format list).
* /documents/extract: multipart upload → {engine, pages, text, tables} with
  per-format engines - pdfplumber (PDF text + ruled tables), pytesseract
  (image OCR), python-docx, openpyxl (sheets = tables), csv/tsv/json/text.
  Unsupported types 415, empty/broken files 400.
* /documents/to-dataset: the best table becomes a first-class v27 dataset
  (source="document") with numeric coercion - qty/amount strings cast to
  real numbers, leading-zero and mixed columns stay text - immediately
  SQL-queryable. Text-only documents 400, dup names 409.
* document_extract node: path source (missing file / unsupported ext → run
  errors), coerced items out, registry grows 36 → 37 visible node types.

Fixtures are generated in-memory (reportlab PDF, PIL PNG, python-docx,
openpyxl) - no binary assets in the repo.
"""

from __future__ import annotations

import asyncio
import io
import os
import uuid

import httpx

from app.main import app

API = "http://testserver/api/v1"
FIXTURE_CSV = "/home/z/my-project/scripts/.v32_node_fixture.csv"


def _client() -> httpx.AsyncClient:
    return httpx.AsyncClient(transport=httpx.ASGITransport(app=app), base_url=API)


async def _drain_background() -> None:
    from app.services import executor as executor_mod

    tasks = [t for t in executor_mod._background_tasks if not t.done()]
    if tasks:
        await asyncio.gather(*tasks, return_exceptions=True)


async def _cleanup(dataset_refs: list[str], workflow_ids: list[str]) -> None:
    async with _client() as client:
        for ref in dataset_refs:
            try:
                await client.delete(f"/datasets/{ref}")
            except Exception:
                pass
        for wid in workflow_ids:
            try:
                await client.delete(f"/workflows/{wid}")
            except Exception:
                pass
    await _drain_background()


def _node(nid: str, ntype: str, params: dict | None = None, name: str | None = None) -> dict:
    return {"id": nid, "type": ntype, "name": name or nid, "position": {"x": 0, "y": 0}, "parameters": params or {}}


def _edge(eid: str, source: str, target: str) -> dict:
    return {"id": eid, "source": source, "target": target, "sourceHandle": "main", "targetHandle": "main"}


async def _make_workflow(client: httpx.AsyncClient, name: str, graph: dict) -> str:
    res = await client.post("/workflows", json={"name": name, "graph": graph, "is_active": False})
    assert res.status_code == 201, res.text
    return res.json()["id"]


async def _run_and_wait(client: httpx.AsyncClient, workflow_id: str, payload: dict | None = None) -> dict:
    res = await client.post(f"/workflows/{workflow_id}/run", json={"payload": payload or {}})
    assert res.status_code in (200, 202), res.text
    exec_id = res.json()["execution_id"]
    for _ in range(100):
        res = await client.get(f"/executions/{exec_id}")
        assert res.status_code == 200, res.text
        if res.json()["status"] != "running":
            return res.json()
        await asyncio.sleep(0.05)
    raise AssertionError("execution did not finish in time")


def _find_node_run(execution: dict, node_name: str) -> dict | None:
    for run in execution.get("node_runs") or []:
        if run.get("node_name") == node_name:
            return run
    return None


# ----------------------------------------------------------------- fixtures
def _invoice_pdf() -> bytes:
    """Ruled invoice table - pdfplumber's lattice detection needs the grid lines."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

    buf = io.BytesIO()
    story = [
        Paragraph("ACME Consulting - Invoice INV-2026-004", getSampleStyleSheet()["Title"]),
        Paragraph("Client: Globex Ltd. Terms: Net 30", getSampleStyleSheet()["Normal"]),
    ]
    rows = [
        ["Item", "Qty", "Unit Price", "Amount"],
        ["Data pipeline build", "1", "4200.00", "4200.00"],
        ["Support retainer", "3", "350.00", "1050.00"],
        ["Training workshop", "2", "500.00", "1000.00"],
    ]
    tbl = Table(rows)
    tbl.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        )
    )
    story.append(tbl)
    story.append(Paragraph("Total Due: 6250.00 USD", getSampleStyleSheet()["Normal"]))
    SimpleDocTemplate(buf, pagesize=letter).build(story)
    return buf.getvalue()


def _text_only_pdf() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=letter).build(
        [Paragraph("A quiet memo with no tables at all. Just prose.", getSampleStyleSheet()["Normal"])]
    )
    return buf.getvalue()


def _ocr_png() -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (640, 200), "white")
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 30)
    except Exception:
        font = ImageFont.load_default()
    d.text((20, 30), "INVOICE TOTAL 6250 USD", fill="black", font=font)
    d.text((20, 100), "DUE DATE 2026-09-30", fill="black", font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# ---------------------------------------------------------------------------
# 1) engines probe + text-format extraction (txt/csv/tsv/json) + guards
# ---------------------------------------------------------------------------
def test_v32_engines_and_text_formats():
    async def _go():
        async with _client() as client:
            res = await client.get("/documents/engines")
            assert res.status_code == 200
            # strict pin lives in the latest wave only (v33 convention)
            hres = await client.get("/health")
            assert hres.status_code == 200 and hres.json()["app"] == "Py8n"
            cap = res.json()
            assert cap["ocr"]["available"] is True and cap["ocr"]["version"]
            assert ".pdf" in cap["formats"] and ".png" in cap["formats"]

            # txt → text engine, no tables
            res = await client.post(
                "/documents/extract",
                files={"file": ("notes.txt", b"quarterly outlook\nsteady growth", "text/plain")},
            )
            assert res.status_code == 200, res.text
            r = res.json()
            assert r["engine"] == "text" and r["pages"] == 1 and r["tables"] == []
            assert "steady growth" in r["text"]

            # csv → one 3-row table
            res = await client.post(
                "/documents/extract",
                files={"file": ("city.csv", b"city,units\nLagos,35\nBerlin,25\n", "text/csv")},
            )
            r = res.json()
            assert r["engine"] == "csv" and r["tables"][0]["n_rows"] == 3
            assert r["tables"][0]["rows"][0] == ["city", "units"]

            # json list-of-dicts → header from union of keys
            res = await client.post(
                "/documents/extract",
                files={"file": ("d.json", b'[{"a": 1, "b": "x"}, {"a": 2, "b": "y"}]', "application/json")},
            )
            r = res.json()
            assert r["engine"] == "json" and r["tables"][0]["n_rows"] == 3
            assert r["tables"][0]["rows"][0] == ["a", "b"]

            # unsupported type 415, empty file 400
            res = await client.post("/documents/extract", files={"file": ("v.exe", b"MZ...", "application/x-msdownload")})
            assert res.status_code == 415 and "Unsupported" in res.json()["detail"]
            res = await client.post("/documents/extract", files={"file": ("empty.txt", b"", "text/plain")})
            assert res.status_code == 400 and "empty" in res.json()["detail"]

    try:
        asyncio.run(_go())
    finally:
        asyncio.run(_cleanup([], []))


# ---------------------------------------------------------------------------
# 2) PDF (pdfplumber text + ruled table) and image OCR
# ---------------------------------------------------------------------------
def test_v32_pdf_and_ocr():
    async def _go():
        async with _client() as client:
            pdf_bytes = _invoice_pdf()
            res = await client.post(
                "/documents/extract",
                files={"file": ("invoice.pdf", pdf_bytes, "application/pdf")},
            )
            assert res.status_code == 200, res.text
            r = res.json()
            assert r["engine"] == "pdf" and r["pages"] == 1
            assert "INV-2026-004" in r["text"] and "Total Due: 6250.00" in r["text"]
            assert len(r["tables"]) == 1
            t = r["tables"][0]
            assert t["rows"][0] == ["Item", "Qty", "Unit Price", "Amount"]
            assert t["n_rows"] == 4 and t["page"] == 1

            # OCR: PIL-generated PNG → tesseract reads both lines
            res = await client.post(
                "/documents/extract",
                files={"file": ("scan.png", _ocr_png(), "image/png")},
            )
            assert res.status_code == 200, res.text
            r = res.json()
            assert r["engine"] == "ocr" and r["pages"] == 1
            flat = " ".join(r["text"].split())
            assert "INVOICE TOTAL 6250 USD" in flat, flat
            assert "DUE DATE 2026-09-30" in flat

    try:
        asyncio.run(_go())
    finally:
        asyncio.run(_cleanup([], []))


# ---------------------------------------------------------------------------
# 3) docx (paragraph + table) and xlsx (sheets = tables)
# ---------------------------------------------------------------------------
def test_v32_docx_and_xlsx():
    async def _go():
        async with _client() as client:
            import docx as docx_lib

            buf = io.BytesIO()
            d = docx_lib.Document()
            d.add_paragraph("Meeting minutes - budget review")
            t = d.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = "dept"
            t.rows[0].cells[1].text = "budget"
            t.rows[1].cells[0].text = "Platform"
            t.rows[1].cells[1].text = "480000"
            d.save(buf)
            res = await client.post("/documents/extract", files={"file": ("minutes.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
            assert res.status_code == 200, res.text
            r = res.json()
            assert r["engine"] == "docx"
            assert "budget review" in r["text"]
            assert r["tables"][0]["rows"] == [["dept", "budget"], ["Platform", "480000"]]

            import openpyxl

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Sales"
            ws.append(["region", "revenue"])
            ws.append(["Lagos", 12000])
            ws.append(["Berlin", 9000])
            ws2 = wb.create_sheet("Costs")
            ws2.append(["region", "cost"])
            ws2.append(["Lagos", 7000])
            xb = io.BytesIO()
            wb.save(xb)
            res = await client.post("/documents/extract", files={"file": ("books.xlsx", xb.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")})
            assert res.status_code == 200, res.text
            r = res.json()
            assert r["engine"] == "xlsx" and r["pages"] == 2
            assert r["meta"]["sheets"] == ["Sales", "Costs"]
            assert len(r["tables"]) == 2
            assert r["tables"][0]["rows"][1] == ["Lagos", "12000"]

    try:
        asyncio.run(_go())
    finally:
        asyncio.run(_cleanup([], []))


# ---------------------------------------------------------------------------
# 4) to-dataset: best table → dataset with coercion → SQL-queryable
# ---------------------------------------------------------------------------
def test_v32_to_dataset_and_sql():
    tag = uuid.uuid4().hex[:8]
    created: list[str] = []

    async def _go():
        async with _client() as client:
            res = await client.post(
                "/documents/to-dataset",
                files={"file": ("invoice.pdf", _invoice_pdf(), "application/pdf")},
                data={"name": f"v32 Invoice Items {tag}", "description": "from document AI"},
            )
            assert res.status_code == 201, res.text
            body = res.json()
            meta = body["dataset"]
            created.append(meta["id"])
            assert meta["source"] == "document" and meta["row_count"] == 3
            assert body["extraction"]["engine"] == "pdf"
            assert body["extraction"]["rows_imported"] == 3
            assert body["extraction"]["tables_count"] == 1

            # numeric columns coerced, text column untouched
            res = await client.get(f"/datasets/{meta['id']}/rows")
            rows = res.json()["rows"]
            assert rows[0]["Item"] == "Data pipeline build"
            assert rows[0]["Qty"] == 1 and rows[1]["Qty"] == 3
            assert rows[0]["Amount"] == 4200.0 and rows[2]["Amount"] == 1000.0
            schema = {c["name"]: c["dtype"] for c in meta["schema_json"]}
            assert schema["Qty"] == "integer" and schema["Amount"] == "number"
            assert schema["Item"] == "text"

            # immediately SQL-queryable (view names fold non-alphanumerics to _)
            res = await client.post("/datasets/query", json={"sql": f'SELECT SUM("Amount") AS total FROM "v32_invoice_items_{tag}"'})
            assert res.status_code == 200, res.text
            assert res.json()["rows"][0]["total"] == 6250.0

            # dup name 409; text-only PDF 400 with the "text only" hint
            res = await client.post(
                "/documents/to-dataset",
                files={"file": ("invoice.pdf", _invoice_pdf(), "application/pdf")},
                data={"name": f"V32 INVOICE ITEMS {tag}"},
            )
            assert res.status_code == 409
            res = await client.post(
                "/documents/to-dataset",
                files={"file": ("memo.pdf", _text_only_pdf(), "application/pdf")},
                data={"name": f"v32 Memo {tag}"},
            )
            assert res.status_code == 400 and "text only" in res.json()["detail"]

            # name defaults from the filename when absent
            res = await client.post(
                "/documents/to-dataset",
                files={"file": (f"v32 default {tag}.csv", b"city,units\nLagos,35\n", "text/csv")},
                data={},
            )
            assert res.status_code == 201, res.text
            assert res.json()["dataset"]["name"] == f"v32 default {tag}"
            created.append(res.json()["dataset"]["id"])

    try:
        asyncio.run(_go())
    finally:
        asyncio.run(_cleanup(created, []))


# ---------------------------------------------------------------------------
# 5) document_extract node: definitions, happy path, run errors
# ---------------------------------------------------------------------------
def test_v32_document_extract_node():
    tag = uuid.uuid4().hex[:8]
    wf_ids: list[str] = []

    async def _go():
        async with _client() as client:
            res = await client.get("/node-definitions")
            types = [d["type"] for d in res.json()["definitions"]]
            assert len(types) >= 37, f"expected 37+ visible types, got {len(types)}"  # 45 at v45
            doc_def = next(d for d in res.json()["definitions"] if d["type"] == "document_extract")
            assert doc_def["icon"] == "file-text" and doc_def["category"] == "actions"
            props = doc_def["parameters_schema"]["properties"]
            assert doc_def["defaults"].get("source") == "path" and "url" in props

            # fixture csv on disk
            with open(FIXTURE_CSV, "wb") as f:
                f.write(b"city,units\nLagos,35\nBerlin,25\n")

            graph = {
                "nodes": [
                    _node("t", "manual_trigger"),
                    _node("d", "document_extract", {"source": "path", "path": FIXTURE_CSV}, "Extract"),
                    _node("s", "set_variable", {"keep_input": False, "assignments": {"n": "{{ input.items_count }}", "first": "{{ input.items[0].city }}"}}, "Meta"),
                ],
                "edges": [_edge("e1", "t", "d"), _edge("e2", "d", "s")],
            }
            wid = await _make_workflow(client, f"v32 extract {tag}", graph)
            wf_ids.append(wid)
            run = await _run_and_wait(client, wid)
            assert run["status"] == "success", run.get("error")

            ext = _find_node_run(run, "Extract")
            assert ext["status"] == "success"
            out = ext["output"]
            assert out["engine"] == "csv" and out["items_count"] == 2
            assert out["items"][0]["city"] == "Lagos" and out["items"][0]["units"] == 35  # coerced
            assert "city,units" in out["text_preview"]

            meta = _find_node_run(run, "Meta")
            assert meta["output"] == {"n": 2, "first": "Lagos"}

            # missing file → run error
            graph2 = {
                "nodes": [
                    _node("t", "manual_trigger"),
                    _node("d", "document_extract", {"source": "path", "path": "/nonexistent/ghost.pdf"}),
                ],
                "edges": [_edge("e1", "t", "d")],
            }
            wid2 = await _make_workflow(client, f"v32 ghost {tag}", graph2)
            wf_ids.append(wid2)
            run2 = await _run_and_wait(client, wid2)
            assert run2["status"] == "error" and "File not found" in (run2.get("error") or "")

            # unsupported extension → run error
            with open(FIXTURE_CSV, "wb") as f:
                f.write(b"binary-ish")
            exe_path = FIXTURE_CSV + ".exe"
            os.replace(FIXTURE_CSV, exe_path)
            graph3 = {
                "nodes": [
                    _node("t", "manual_trigger"),
                    _node("d", "document_extract", {"source": "path", "path": exe_path}),
                ],
                "edges": [_edge("e1", "t", "d")],
            }
            wid3 = await _make_workflow(client, f"v32 exe {tag}", graph3)
            wf_ids.append(wid3)
            run3 = await _run_and_wait(client, wid3)
            assert run3["status"] == "error" and "Unsupported document type" in (run3.get("error") or "")

    try:
        asyncio.run(_go())
    finally:
        if os.path.exists(FIXTURE_CSV):
            os.remove(FIXTURE_CSV)
        if os.path.exists(FIXTURE_CSV + ".exe"):
            os.remove(FIXTURE_CSV + ".exe")
        asyncio.run(_cleanup([], wf_ids))


# ---------------------------------------------------------------------------
# 6) coercion helper edge cases (unit-level)
# ---------------------------------------------------------------------------
def test_v32_coercion_rules():
    from app.services import documents as doc_svc

    # ints, floats, leading zeros stay text, mixed column stays text, blanks tolerated
    items = [
        {"qty": "1", "amount": "4200.00", "code": "007", "note": "x"},
        {"qty": "3", "amount": "1050.00", "code": "042", "note": "5"},
        {"qty": "2", "amount": "", "code": "009", "note": ""},
    ]
    out = doc_svc.coerce_items_dtypes(items)
    assert out[0]["qty"] == 1 and out[1]["qty"] == 3 and isinstance(out[0]["qty"], int)
    assert out[0]["amount"] == 4200.0 and out[2]["amount"] == ""  # blank kept
    assert out[0]["code"] == "007" and out[0]["note"] == "x"  # leading zeros + mixed stay text

    # negatives (ints) and dot-leading floats - in SEPARATE columns (a column
    # mixing ints and floats stays text by design)
    out2 = doc_svc.coerce_items_dtypes([{"a": "-12", "b": ".5"}, {"a": "3", "b": "2.5"}])
    assert out2[0]["a"] == -12 and isinstance(out2[0]["a"], int)
    assert out2[1]["b"] == 2.5 and isinstance(out2[1]["b"], float)

    # mixed int/float single column stays text
    out3 = doc_svc.coerce_items_dtypes([{"v": "-12"}, {"v": "3"}, {"v": ".5"}])
    assert out3[0]["v"] == "-12" and out3[2]["v"] == ".5"

    # header handling: empty header cell → col_N, duplicate → _2
    table = {"page": 1, "rows": [["city", "", "city"], ["Lagos", "1", "2"]], "n_rows": 2, "n_cols": 3}
    items = doc_svc.table_to_items(table)
    assert list(items[0].keys()) == ["city", "col_2", "city_2"]
    assert items[0]["city_2"] == "2"

    # best_table picks the row-heaviest table
    best = doc_svc.best_table(
        {
            "tables": [
                {"page": 1, "rows": [["a"], ["1"]], "n_rows": 2},
                {"page": 2, "rows": [["b"], ["1"], ["2"], ["3"]], "n_rows": 4},
            ]
        }
    )
    assert best["page"] == 2 and best["n_rows"] == 4
    assert doc_svc.best_table({"tables": [{"page": 1, "rows": [["just", "header"]], "n_rows": 1}]}) is None


# NOTE: old node count assertions live in test_v10 (websocket), test_v22 (355),
# test_v24 (420), test_v25 (95), test_v27 (100), test_v28 (129), test_v30 (124),
# test_v31 (85) - all bumped 36 → 37 by scripts/bump_counts_v32.py alongside
# the smoke test.
